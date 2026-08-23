"""Update docx: paragraphs [13] and [14] only. Everything else untouched."""
from docx import Document
from docx.shared import RGBColor
import os

path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                     'research_proposal_text.docx')
doc = Document(path)

GRAY = RGBColor(0x88, 0x88, 0x88)

# Paragraph [13] — the "train-then-optimize" paragraph
P13_START = "The common approach to this problem"
P13_NEW = (
    "The common approach to this problem is to separate the training into two "
    "stages: (1) train a classifier to maximize accuracy, then (2) apply "
    "post-hoc adjustments such as top-K selection or greedy reallocation to "
    "enforce prediction counts. This two-stage paradigm, commonly referred to "
    "as predict-then-optimize [REF: elmachtoub2022smart], is limited because "
    "the model is never exposed to the constrained structure of the problem "
    "during training. Its decision boundaries are not shaped to reach "
    "constraint satisfaction."
)

# Paragraph [14] — the "Several lines of work" paragraph
P14_START = "Several lines of work address related problems"
P14_NEW = (
    "Several lines of work address related problems. Cost-sensitive learning "
    "[REF: elkan2001foundations] assigns different misclassification costs to "
    "different classes, and loss-function modifications such as focal loss "
    "[REF: lin2017focal] and class-balanced loss [REF: cui2019class] reshape "
    "training objectives to address class imbalance, but none enforce hard "
    "prediction-count limits. Resource allocation frameworks "
    "[REF: shifman2023adaptive, cohen2024resource, vanderschueren2024perspective] "
    "optimize class assignments subject to capacity constraints using "
    "mathematical programming, but operate post-hoc on a fixed classifier "
    "rather than during training. Fairness-aware classification "
    "[REF: hardt2016equality] enforces constraints across subgroups, but "
    "regulates prediction rates rather than absolute counts \u2014 an important "
    "distinction in resource-allocation settings where a fixed number of "
    "actions (e.g., biopsies) is available. Constrained optimization has been "
    "applied to neural network training for class-imbalance objectives "
    "[REF: sangalli2021constrained] and with theoretical guarantees for "
    "non-convex losses [REF: chamon2023constrained, hounie2023resilient]; "
    "indeed, recent work advocates for broader adoption of constrained methods "
    "over fixed penalties in deep learning [REF: ramirez2025position]. "
    "However, these approaches address rate-based metrics rather than "
    "prediction-count budgets. To date, no approach offers a mechanism for "
    "enforcing multi-level budgeted constraints during neural network training."
)

updates = [
    (P13_START, P13_NEW),
    (P14_START, P14_NEW),
]

for start_text, new_text in updates:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(start_text):
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)
            print(f"Updated paragraph [{i}] starting with '{start_text[:40]}...'")
            break

doc.save(path)
print(f"Saved {path}")
